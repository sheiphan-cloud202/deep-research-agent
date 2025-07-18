"""
Citation Report Generator Agent - PDF Strategy Reports with Clickable Links

This agent generates a single comprehensive PDF business strategy report with clickable
citation links that consolidates all use cases from the ideation phase. Key capabilities:

1. PDF-FIRST APPROACH: Primary output is a professional PDF with clickable citation links
2. CONSOLIDATED REPORTING: Creates one unified report instead of individual reports
3. STAGED IMPLEMENTATION: Organizes use cases by priority/complexity into 3 phases:
   - Phase 1: Immediate Wins (0-6 months) - High priority, low-medium complexity
   - Phase 2: Strategic Expansion (6-18 months) - Medium priority or high complexity
   - Phase 3: Strategic Innovation (18+ months) - Long-term strategic initiatives

4. CLICKABLE CITATIONS: All citations are hyperlinked and clickable in the final PDF
5. COMPREHENSIVE CITATION SEARCH: Searches across all use cases to find relevant citations
6. MULTI-FORMAT OUTPUT: Primary PDF + Secondary formats (Markdown, HTML, DOCX)
7. S3 INTEGRATION: Optionally uploads consolidated report files to AWS S3

REPORT STRUCTURE:
- Executive Summary with strategic value and business impact
- Market Context & Strategic Rationale with citations
- 3-Phase Implementation Roadmap with use cases organized by stage
- Technology Stack & AI Applications analysis
- Business Impact Analysis with quantified projections
- Risk Assessment & Mitigation strategies
- Success Metrics & KPIs for each phase
- Conclusion & Next Steps

Dependencies (Pure Python - No External Tools Required):
- SERPER_API_KEY environment variable for citation search
- markdown: Python markdown to HTML conversion (pip install markdown)
- xhtml2pdf: Pure Python HTML to PDF with clickable links (pip install xhtml2pdf)
- python-docx: DOCX generation (pip install python-docx)
- boto3: For S3 upload functionality (optional)

INSTALLATION (No System Dependencies Required):
  Python packages only: pip install markdown xhtml2pdf python-docx
  All packages are pure Python - no external tools or system libraries needed!

Environment Variables:
- SERPER_API_KEY: API key for Serper.dev search service
- S3_BUCKET: S3 bucket name for file uploads (optional)
- AWS credentials for S3 upload (optional)

The agent gracefully handles missing dependencies and will continue to function
with reduced capabilities if optional tools are not available.
"""

import os
import re
from typing import Any

import boto3
import markdown
import requests
from dotenv import load_dotenv

from deep_research_agent.agents.base_agent import BaseAgent
from deep_research_agent.common.schemas import AgentType
from deep_research_agent.core.agent_factory import AgentFactory
from deep_research_agent.services.prompt_service import PromptService
from deep_research_agent.utils.logger import logger

# Load .env variables
load_dotenv()


class CitationReportGeneratorAgent(BaseAgent):
    """
    Agent that generates comprehensive business reports with citations.
    Searches for citations, generates markdown reports, converts to multiple formats,
    and uploads to S3.
    """

    def __init__(self, prompt_service: PromptService, model_id: str | None = None):
        super().__init__(prompt_service)
        self._agent = AgentFactory.create_agent(model_id)
        if self.prompt_service:
            self._agent.system_prompt = self.prompt_service.get_system_prompt(AgentType.CITATION_REPORT_GENERATOR)

        # Configuration
        self.serper_api_key = os.getenv("SERPER_API_KEY")
        self.s3_bucket = os.getenv("S3_BUCKET", "deepresearch-qubitz-document-bucket")

    def execute(self, context: dict[str, Any]):
        """Execute the citation report generation process"""
        try:
            # Get use cases from context (prioritize refined_ideas over initial_ideas)
            use_cases = context.get("refined_ideas") or context.get("initial_ideas")
            if not use_cases:
                logger.warning("No use cases found in context for report generation")
                context["citation_report"] = None
                return

            # Extract use cases list
            use_cases_list = use_cases.use_cases if hasattr(use_cases, "use_cases") else use_cases
            if not use_cases_list:
                logger.warning("No use cases found in the use_cases object")
                context["citation_report"] = None
                return

            # Generate one consolidated report for all use cases
            report_data = self._generate_consolidated_report(use_cases_list, context)

            context["citation_report"] = report_data
            logger.info(f"Generated consolidated citation report covering {len(use_cases_list)} use cases")

        except Exception as e:
            logger.error(f"Citation report generation failed: {e}")
            context["citation_report"] = None
            raise

    def _generate_consolidated_report(self, use_cases_list: list, context: dict[str, Any]) -> dict[str, Any]:
        """Generate a comprehensive consolidated report for all use cases"""

        project_id = context.get("project_id", "consolidated_report")
        user_id = context.get("user_id", "default_user")
        conversation_id = context.get("conversation_id", f"conv_{project_id}")

        logger.info(f"Generating consolidated report for {len(use_cases_list)} use cases")

        # Step 1: Search for citations covering all use cases
        citations = self._search_consolidated_citations(use_cases_list)

        # Step 2: Organize use cases by priority/complexity for staged reporting
        organized_use_cases = self._organize_use_cases_by_stages(use_cases_list)

        # Step 3: Generate consolidated markdown report
        markdown_content = self._generate_consolidated_markdown_report(organized_use_cases, citations, context)

        # Step 4: Save and convert files
        file_paths = self._save_and_convert_files(markdown_content, "consolidated_business_strategy_report")

        # Step 5: Upload to S3 (optional, only if credentials available)
        upload_urls = self._upload_to_s3(file_paths, user_id, conversation_id)

        # Get primary PDF output path
        primary_pdf = file_paths.get("pdf", file_paths.get("primary_output"))

        return {
            "report_type": "consolidated_pdf_strategy_report",
            "use_cases_count": len(use_cases_list),
            "organized_use_cases": organized_use_cases,
            "markdown_content": markdown_content,
            "file_paths": file_paths,
            "primary_output": primary_pdf,
            "pdf_with_clickable_links": primary_pdf,
            "upload_urls": upload_urls,
            "citations_count": len(citations),
            "status": "PDF report with clickable citation links generated successfully"
            if primary_pdf
            else "Report generated (PDF conversion may have failed)",
        }

    def _extract_keywords(self, text: str, max_words: int = 50) -> str:
        """Extract keywords from text for search queries"""
        words = re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
        top_keywords = list(dict.fromkeys(words))[:max_words]
        return " ".join(top_keywords)

    def _search_consolidated_citations(
        self, use_cases_list: list, max_results: int = 15
    ) -> list[tuple[str, str, str]]:
        """Search for citations covering all use cases using Serper API"""
        if not self.serper_api_key:
            logger.warning("SERPER_API_KEY not found, skipping citation search")
            return []

        try:
            # Extract keywords from all use cases to create comprehensive search queries
            all_keywords = []
            for use_case in use_cases_list:
                if hasattr(use_case, "title"):
                    title = use_case.title
                    description = use_case.description
                else:
                    title = use_case.get("title", "")
                    description = use_case.get("description", "")

                keywords = self._extract_keywords(f"{title} {description}", max_words=10)
                all_keywords.extend(keywords.split())

            # Create consolidated search queries focusing on AI, business strategy, and implementation
            search_queries = [
                " ".join(list(dict.fromkeys(all_keywords))[:15]) + " AI business strategy",
                " ".join(list(dict.fromkeys(all_keywords))[:15]) + " market analysis implementation",
                " ".join(list(dict.fromkeys(all_keywords))[:10]) + " technology trends competitive advantage",
            ]

            all_citations = []
            seen = set()

            for query in search_queries:
                headers = {"X-API-KEY": self.serper_api_key, "Content-Type": "application/json"}
                payload = {"q": query, "gl": "us", "hl": "en", "num": max_results // len(search_queries) + 2}

                logger.info(f"Searching citations for: {query}")

                response = requests.post("https://google.serper.dev/search", headers=headers, json=payload)
                response.raise_for_status()

                serper_results = response.json().get("organic", [])

                # Process results and remove duplicates
                for result in serper_results:
                    link = result.get("link") or result.get("url")
                    title = result.get("title", "Untitled")
                    snippet = result.get("snippet", "")

                    if link and link not in seen and len(all_citations) < max_results:
                        seen.add(link)
                        all_citations.append((title, link, snippet))

            logger.info(f"Found {len(all_citations)} unique citations across all use cases")
            return all_citations

        except Exception as e:
            logger.error(f"Consolidated citation search failed: {e}")
            return []

    def _organize_use_cases_by_stages(self, use_cases_list: list) -> dict[str, list]:
        """Organize use cases by implementation stages based on priority and complexity"""
        stages = {
            "immediate": [],  # High priority, low-medium complexity
            "phase_2": [],  # Medium priority or high complexity
            "strategic": [],  # Long-term strategic initiatives
        }

        for use_case in use_cases_list:
            if hasattr(use_case, "priority"):
                priority = getattr(use_case, "priority", "Medium").lower()
                complexity = getattr(use_case, "complexity", "Medium").lower()
            else:
                priority = use_case.get("priority", "medium").lower()
                complexity = use_case.get("complexity", "medium").lower()

            # Staging logic
            if priority in ["critical", "high"] and complexity in ["low", "medium"]:
                stages["immediate"].append(use_case)
            elif priority == "medium" or complexity == "high":
                stages["phase_2"].append(use_case)
            else:
                stages["strategic"].append(use_case)

        # Ensure balanced distribution - move items if one category is empty
        if not stages["immediate"] and stages["phase_2"]:
            stages["immediate"].append(stages["phase_2"].pop(0))
        if not stages["phase_2"] and stages["strategic"]:
            stages["phase_2"].append(stages["strategic"].pop(0))

        logger.info(
            f"Organized use cases: {len(stages['immediate'])} immediate, "
            f"{len(stages['phase_2'])} phase 2, {len(stages['strategic'])} strategic"
        )

        return stages

    def _generate_consolidated_markdown_report(
        self, organized_use_cases: dict[str, list], citations: list[tuple[str, str, str]], context: dict[str, Any]
    ) -> str:
        """Generate consolidated markdown report using the LLM agent"""

        # Format citation block
        citation_block = "\n".join(
            [f"[[{i + 1}]] [{title}]({link}) — {snippet}" for i, (title, link, snippet) in enumerate(citations)]
        )

        if not self.prompt_service:
            raise ValueError("PromptService is not available for CitationReportGeneratorAgent")

        # Get additional context data
        research_summary = context.get("research_summary", "")
        conversation_summary = context.get("conversation_summary", "")

        # Format use cases by stage for the prompt
        use_cases_by_stage = {}
        for stage, use_cases in organized_use_cases.items():
            formatted_cases = []
            for i, use_case in enumerate(use_cases):
                if hasattr(use_case, "title"):
                    title = use_case.title
                    description = use_case.description
                    business_value = getattr(use_case, "business_value", "")
                else:
                    title = use_case.get("title", f"Use Case {i + 1}")
                    description = use_case.get("description", "")
                    business_value = use_case.get("business_value", "")

                formatted_cases.append({"title": title, "description": description, "business_value": business_value})
            use_cases_by_stage[stage] = formatted_cases

        prompt = self.prompt_service.format_user_prompt(
            AgentType.CITATION_REPORT_GENERATOR,
            "generate_consolidated_report",
            use_cases_by_stage=str(use_cases_by_stage),
            citation_block=citation_block,
            research_summary=research_summary,
            conversation_summary=conversation_summary,
            total_use_cases=sum(len(cases) for cases in organized_use_cases.values()),
        )

        result = self._agent(prompt)
        return str(result).strip()

    def _save_and_convert_files(self, markdown_content: str, use_case_id: str) -> dict[str, str]:
        """Save markdown and convert to PDF with clickable links using pure Python packages"""
        file_paths = {}

        try:
            # Save markdown file
            md_filename = f"{use_case_id}_report.md"
            with open(md_filename, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            file_paths["markdown"] = md_filename
            logger.info(f"Saved markdown: {md_filename}")

            # Convert markdown to HTML using Python markdown package
            try:
                html_path = md_filename.replace(".md", ".html")

                # Pre-process markdown to fix citation links for better PDF conversion
                processed_markdown = self._preprocess_citation_links(markdown_content)

                # Configure markdown with extensions for better formatting and link handling
                md = markdown.Markdown(extensions=["tables", "fenced_code", "toc", "attr_list"])

                # Convert processed markdown to HTML
                html_content = md.convert(processed_markdown)

                # Create a complete HTML document with CSS for better PDF rendering
                full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Business Strategy Report</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            line-height: 1.2;
            max-width: 210mm;
            margin: 0 auto;
            padding: 15mm;
            color: #333;
            font-size: 11pt;
        }}
        p {{
            margin: 2pt 0;
            text-align: justify;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 8pt;
            font-size: 24pt;
            page-break-before: always;
            margin: 0 0 12pt 0;
        }}
        h1:first-child {{
            page-break-before: avoid;
        }}
        h2 {{
            color: #34495e;
            border-bottom: 1px solid #bdc3c7;
            padding-bottom: 4pt;
            margin: 12pt 0 8pt 0;
            font-size: 18pt;
            page-break-before: always;
        }}
        h2.no-break {{
            page-break-before: avoid;
        }}
        h3 {{
            color: #2980b9;
            margin: 8pt 0 4pt 0;
            font-size: 14pt;
        }}
        h4 {{
            color: #7f8c8d;
            margin: 6pt 0 3pt 0;
            font-size: 12pt;
        }}
        a {{
            color: #2980b9 !important;
            text-decoration: none !important;
            font-weight: bold !important;
            border-bottom: 1px solid #2980b9 !important;
        }}
        a:hover {{
            background-color: #e8f4fd !important;
        }}
        .citation-link {{
            color: #2980b9 !important;
            text-decoration: none !important;
            font-weight: bold !important;
            border-bottom: 1px solid #2980b9 !important;
            padding: 1px 2px !important;
            background-color: #f8f9fa !important;
        }}
        strong {{
            color: #c0392b;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 8pt 0;
        }}
        th, td {{
            border: 1px solid #bdc3c7;
            padding: 6pt;
            text-align: left;
            font-size: 10pt;
        }}
        th {{
            background-color: #ecf0f1;
            font-weight: bold;
        }}
        ul, ol {{
            margin: 4pt 0;
            padding-left: 20pt;
        }}
        li {{
            margin: 2pt 0;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 8pt 0;
            padding: 6pt 12pt;
            background-color: #f8f9fa;
        }}
        @page {{
            size: A4;
            margin: 15mm;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

                # Save HTML file
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(full_html)
                file_paths["html"] = html_path
                logger.info(f"HTML created with enhanced styling: {html_path}")

                # TIER 1: Try WeasyPrint for excellent HTML/CSS support (moved to primary)
                pdf_path = md_filename.replace(".md", ".pdf")
                pdf_generated = False

                try:
                    pdf_generated = self._generate_pdf_weasyprint(full_html, pdf_path)
                    if pdf_generated:
                        file_paths["pdf"] = pdf_path
                        file_paths["primary_output"] = pdf_path
                        logger.info(f"✅ High-quality PDF with clickable links created using WeasyPrint: {pdf_path}")
                except ImportError:
                    logger.info("WeasyPrint not available, trying pypandoc...")
                except Exception as e:
                    logger.warning(f"WeasyPrint generation failed: {e}, trying pypandoc...")

                # TIER 2: Try pypandoc for excellent markdown handling
                if not pdf_generated:
                    try:
                        pdf_generated = self._generate_pdf_pypandoc(processed_markdown, pdf_path)
                        if pdf_generated:
                            file_paths["pdf"] = pdf_path
                            file_paths["primary_output"] = pdf_path
                            logger.info(f"✅ Professional PDF created using pypandoc: {pdf_path}")
                    except ImportError:
                        logger.info("pypandoc not available, trying ReportLab...")
                    except Exception as e:
                        logger.warning(f"pypandoc generation failed: {e}, trying ReportLab...")

                # TIER 3: Try ReportLab with improved link handling
                if not pdf_generated:
                    try:
                        pdf_generated = self._generate_pdf_reportlab_improved(processed_markdown, pdf_path)
                        if pdf_generated:
                            file_paths["pdf"] = pdf_path
                            file_paths["primary_output"] = pdf_path
                            logger.info(f"✅ Professional PDF created using ReportLab: {pdf_path}")
                    except ImportError:
                        logger.info("ReportLab not available, trying xhtml2pdf...")
                    except Exception as e:
                        logger.warning(f"ReportLab generation failed: {e}, trying xhtml2pdf...")

                # TIER 4: Fallback to enhanced xhtml2pdf
                if not pdf_generated:
                    try:
                        pdf_generated = self._generate_pdf_xhtml2pdf(full_html, pdf_path)
                        if pdf_generated:
                            file_paths["pdf"] = pdf_path
                            file_paths["primary_output"] = pdf_path
                            logger.info(f"✅ Basic PDF created using xhtml2pdf: {pdf_path}")
                    except Exception as e:
                        logger.warning(f"xhtml2pdf generation failed: {e}, creating HTML export...")

                # TIER 5: Emergency HTML export if all PDF generation fails
                if not pdf_generated:
                    logger.warning("All PDF generation methods failed, providing HTML export")
                    file_paths["primary_output"] = html_path
                    logger.info(f"📄 HTML export available: {html_path}")

            except Exception as e:
                logger.warning(f"HTML conversion failed: {e}")

            # Create DOCX as secondary output using a simple approach
            try:
                from docx import Document
                from docx.shared import Inches

                docx_path = md_filename.replace(".md", ".docx")
                doc = Document()

                # Add title
                title = doc.add_heading("AI-Powered Business Strategy Report", 0)

                # Add content (simplified - just the text without full formatting)
                # This is a basic implementation - for full markdown to DOCX conversion
                # you would need a more sophisticated parser
                content_para = doc.add_paragraph()
                content_para.add_run("This is a simplified DOCX version. ")
                content_para.add_run(
                    "For the complete formatted report with clickable links, please refer to the PDF version."
                )

                # Add note about markdown content
                doc.add_paragraph("Original markdown content:")
                doc.add_paragraph(
                    markdown_content[:1000] + "..." if len(markdown_content) > 1000 else markdown_content
                )

                doc.save(docx_path)
                file_paths["docx"] = docx_path
                logger.info(f"Basic DOCX created: {docx_path}")

            except ImportError:
                logger.warning("python-docx not available for DOCX generation")
            except Exception as e:
                logger.warning(f"DOCX generation failed: {e}")

        except Exception as e:
            logger.error(f"File conversion failed: {e}")

        return file_paths

    def _preprocess_citation_links(self, markdown_content: str) -> str:
        """Pre-process markdown to fix citation links for better PDF conversion"""
        # Convert [[15]](URL) format to [15](URL) format for better link handling
        processed = re.sub(r"\[\[(\d+)\]\]\(([^)]+)\)", r"[\1](\2)", markdown_content)

        # Also handle any other double bracket citations that might not have been caught
        processed = re.sub(r"\[\[([^\]]+)\]\]\(([^)]+)\)", r"[\1](\2)", processed)

        return processed

    def _generate_pdf_pypandoc(self, markdown_content: str, pdf_path: str) -> bool:
        """TIER 2: Generate PDF using pypandoc for excellent markdown handling"""
        try:
            import pypandoc  # type: ignore

            # Create a temporary markdown file
            temp_md = pdf_path.replace(".pdf", "_temp.md")
            with open(temp_md, "w", encoding="utf-8") as f:
                f.write(markdown_content)

            # Convert markdown to PDF using pypandoc
            pypandoc.convert_file(
                temp_md,
                "pdf",
                outputfile=pdf_path,
                extra_args=[
                    "--pdf-engine=xelatex",
                    "--variable",
                    "geometry:margin=0.75in",
                    "--variable",
                    "fontsize=11pt",
                    "--variable",
                    "linestretch=1.2",
                    "--variable",
                    "mainfont=Times New Roman",
                    "--toc",
                ],
            )

            # Clean up temp file
            if os.path.exists(temp_md):
                os.remove(temp_md)

            return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0

        except ImportError:
            raise ImportError("pypandoc not available")
        except Exception as e:
            logger.error(f"pypandoc PDF generation error: {e}")
            # Clean up temp file if it exists
            temp_md = pdf_path.replace(".pdf", "_temp.md")
            if os.path.exists(temp_md):
                try:
                    os.remove(temp_md)
                except:
                    pass
            return False

    def _generate_pdf_reportlab_improved(self, markdown_content: str, pdf_path: str) -> bool:
        """TIER 3: Generate PDF using ReportLab with improved link handling"""
        try:
            import re

            from reportlab.lib.colors import HexColor
            from reportlab.lib.enums import TA_JUSTIFY
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate

            # Create professional styles with precise spacing
            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                "ProfessionalTitle",
                parent=styles["Heading1"],
                fontSize=24,
                spaceAfter=6,
                spaceBefore=0,
                textColor=HexColor("#2c3e50"),
                keepWithNext=True,
                fontName="Times-Bold",
            )

            heading_style = ParagraphStyle(
                "ProfessionalHeading",
                parent=styles["Heading2"],
                fontSize=16,
                spaceAfter=4,
                spaceBefore=8,
                textColor=HexColor("#34495e"),
                keepWithNext=True,
                fontName="Times-Bold",
            )

            subheading_style = ParagraphStyle(
                "ProfessionalSubHeading",
                parent=styles["Heading3"],
                fontSize=14,
                spaceAfter=3,
                spaceBefore=6,
                textColor=HexColor("#2980b9"),
                keepWithNext=True,
                fontName="Times-Bold",
            )

            body_style = ParagraphStyle(
                "ProfessionalBody",
                parent=styles["Normal"],
                fontSize=11,
                spaceAfter=3,
                spaceBefore=1,
                alignment=TA_JUSTIFY,
                textColor=HexColor("#333333"),
                fontName="Times-Roman",
                lineHeight=13,
            )

            # Create document with professional margins
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch,
            )

            story = []
            lines = markdown_content.split("\n")
            first_heading = True

            def format_line_improved(text):
                """Improved formatting for citation links"""
                # Handle citation links like [15](URL) - make them clearly clickable
                text = re.sub(
                    r"\[(\d+)\]\(([^)]+)\)", r'<a href="\2"><font color="#2980b9"><u>[\1]</u></font></a>', text
                )
                # Handle regular links [text](URL)
                text = re.sub(
                    r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2"><font color="#2980b9"><u>\1</u></font></a>', text
                )
                # Bold and Italic
                text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
                text = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", text)
                return text

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Handle headings with smart page breaks
                if line.startswith("# "):
                    if not first_heading:
                        story.append(PageBreak())
                    text = format_line_improved(line[2:].strip())
                    story.append(Paragraph(text, title_style))
                    first_heading = False

                elif line.startswith("## "):
                    if story:
                        story.append(PageBreak())
                    text = format_line_improved(line[3:].strip())
                    story.append(Paragraph(text, heading_style))

                elif line.startswith("### "):
                    text = format_line_improved(line[4:].strip())
                    story.append(Paragraph(text, subheading_style))

                elif line.startswith("- ") or line.startswith("* "):
                    text = format_line_improved(line[2:].strip())
                    story.append(Paragraph(f"• {text}", body_style))

                elif line:
                    # Regular paragraph with link and formatting conversion
                    text = format_line_improved(line)
                    story.append(Paragraph(text, body_style))

            # Build PDF
            doc.build(story)

            return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0

        except ImportError:
            raise ImportError("ReportLab not available")
        except Exception as e:
            logger.error(f"ReportLab improved PDF generation error: {e}")
            return False

    def _generate_pdf_reportlab(self, markdown_content: str, pdf_path: str) -> bool:
        """TIER 1: Generate PDF using ReportLab for maximum control"""
        try:
            import re

            from reportlab.lib.colors import HexColor
            from reportlab.lib.enums import TA_JUSTIFY
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate

            # Create professional styles with precise spacing
            styles = getSampleStyleSheet()

            # Allow embedded HTML tags for links, bold, and italic
            styles["Normal"].wordWrap = "CJK"
            styles["BodyText"].wordWrap = "CJK"

            title_style = ParagraphStyle(
                "ProfessionalTitle",
                parent=styles["Heading1"],
                fontSize=24,
                spaceAfter=6,
                spaceBefore=0,
                textColor=HexColor("#2c3e50"),
                keepWithNext=True,
                fontName="Times-Bold",
            )

            heading_style = ParagraphStyle(
                "ProfessionalHeading",
                parent=styles["Heading2"],
                fontSize=16,
                spaceAfter=4,
                spaceBefore=8,
                textColor=HexColor("#34495e"),
                keepWithNext=True,
                fontName="Times-Bold",
            )

            subheading_style = ParagraphStyle(
                "ProfessionalSubHeading",
                parent=styles["Heading3"],
                fontSize=14,
                spaceAfter=3,
                spaceBefore=6,
                textColor=HexColor("#2980b9"),
                keepWithNext=True,
                fontName="Times-Bold",
            )

            body_style = ParagraphStyle(
                "ProfessionalBody",
                parent=styles["Normal"],
                fontSize=11,
                spaceAfter=3,
                spaceBefore=1,
                alignment=TA_JUSTIFY,
                textColor=HexColor("#333333"),
                fontName="Times-Roman",
                lineHeight=13,
            )

            # Create document with professional margins
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch,
            )

            story = []
            lines = markdown_content.split("\n")
            first_heading = True

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Convert markdown formatting to ReportLab HTML tags
                def format_line(text):
                    # Links with explicit color and underline for clickability
                    text = re.sub(
                        r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2"><font color="#2980b9"><u>\1</u></font></a>', text
                    )
                    # Bold and Italic
                    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
                    text = re.sub(r"\*([^*]+)\*", r"<i>\1</i>", text)
                    return text

                # Handle headings with smart page breaks
                if line.startswith("# "):
                    if not first_heading:
                        story.append(PageBreak())
                    text = format_line(line[2:].strip())
                    story.append(Paragraph(text, title_style))
                    first_heading = False

                elif line.startswith("## "):
                    if story:
                        story.append(PageBreak())
                    text = format_line(line[3:].strip())
                    story.append(Paragraph(text, heading_style))

                elif line.startswith("### "):
                    text = format_line(line[4:].strip())
                    story.append(Paragraph(text, subheading_style))

                elif line.startswith("- ") or line.startswith("* "):
                    text = format_line(line[2:].strip())
                    story.append(Paragraph(f"• {text}", body_style))

                elif line:
                    # Regular paragraph with link and formatting conversion
                    text = format_line(line)
                    story.append(Paragraph(text, body_style))

            # Build PDF
            doc.build(story)

            return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0

        except ImportError:
            raise ImportError("ReportLab not available")
        except Exception as e:
            logger.error(f"ReportLab PDF generation error: {e}")
            return False

    def _generate_pdf_weasyprint(self, html_content: str, pdf_path: str) -> bool:
        """TIER 2: Generate PDF using WeasyPrint for excellent HTML/CSS support"""
        try:
            from weasyprint import CSS, HTML

            # Enhanced CSS for professional output
            professional_css = CSS(
                string="""
                @page {
                    size: A4;
                    margin: 0.75in;
                }
                body {
                    font-family: 'Times New Roman', 'Liberation Serif', serif;
                    font-size: 11pt;
                    line-height: 1.2;
                    color: #333;
                }
                h1 {
                    font-size: 24pt;
                    color: #2c3e50;
                    margin: 0 0 6pt 0;
                    page-break-before: always;
                }
                h1:first-child {
                    page-break-before: avoid;
                }
                h2 {
                    font-size: 16pt;
                    color: #34495e;
                    margin: 8pt 0 4pt 0;
                    page-break-before: always;
                }
                h3 {
                    font-size: 14pt;
                    color: #2980b9;
                    margin: 6pt 0 3pt 0;
                }
                p {
                    margin: 1pt 0 3pt 0;
                    text-align: justify;
                }
                a {
                    color: #2980b9;
                    text-decoration: none;
                    font-weight: 500;
                }
                ul, ol {
                    margin: 3pt 0;
                    padding-left: 20pt;
                }
                li {
                    margin: 1pt 0;
                }
            """
            )

            html_doc = HTML(string=html_content)
            pdf_bytes = html_doc.write_pdf(stylesheets=[professional_css])

            if pdf_bytes:
                with open(pdf_path, "wb") as f:
                    f.write(pdf_bytes)
            else:
                return False

            return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0

        except ImportError:
            raise ImportError("WeasyPrint not available")
        except Exception as e:
            logger.error(f"WeasyPrint PDF generation error: {e}")
            return False

    def _generate_pdf_xhtml2pdf(self, html_content: str, pdf_path: str) -> bool:
        """TIER 3: Generate PDF using xhtml2pdf with enhanced CSS"""
        try:
            from xhtml2pdf import pisa

            # Apply aggressive CSS optimization for tighter spacing
            tight_html = (
                html_content.replace("margin: 8px 0;", "margin: 1px 0; line-height: 1.1;")
                .replace("margin-top: 25px;", "margin-top: 4px;")
                .replace("margin-bottom: 15px;", "margin-bottom: 2px;")
                .replace("margin-top: 20px;", "margin-top: 3px;")
                .replace("margin-bottom: 10px;", "margin-bottom: 1px;")
                .replace("margin: 12px 0;", "margin: 2px 0;")
                .replace("margin: 5px 0;", "margin: 1px 0;")
            )

            with open(pdf_path, "w+b") as result_file:
                pisa_status = pisa.CreatePDF(tight_html, dest=result_file, encoding="utf-8")

            return os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0

        except ImportError:
            raise ImportError("xhtml2pdf not available")
        except Exception as e:
            logger.error(f"xhtml2pdf PDF generation error: {e}")
            return False

    def _upload_to_s3(self, file_paths: dict[str, str], user_id: str, conversation_id: str) -> dict[str, str]:
        """Upload files to S3 organized by conversation ID and return URLs"""
        upload_urls = {}

        try:
            s3_client = boto3.client("s3")
            # Organize files by conversation ID: conversations/{conversation_id}/reports/
            s3_prefix = f"conversations/{conversation_id}/reports/"

            logger.info(f"Uploading files to S3 bucket: {self.s3_bucket}")
            logger.info(f"S3 prefix: {s3_prefix}")

            for format_type, local_path in file_paths.items():
                if os.path.exists(local_path):
                    try:
                        # Create descriptive filename with timestamp
                        from datetime import datetime

                        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")

                        file_extension = os.path.splitext(local_path)[1]
                        filename = f"business_strategy_report_{timestamp}{file_extension}"

                        key = s3_prefix + filename

                        # Upload with appropriate content type (no ACL - use bucket default permissions)
                        extra_args = {}

                        # Set content type based on file extension
                        if file_extension == ".pdf":
                            extra_args["ContentType"] = "application/pdf"
                        elif file_extension == ".html":
                            extra_args["ContentType"] = "text/html"
                        elif file_extension == ".md":
                            extra_args["ContentType"] = "text/markdown"
                        elif file_extension == ".docx":
                            extra_args["ContentType"] = (
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                        s3_client.upload_file(local_path, self.s3_bucket, key, ExtraArgs=extra_args)

                        # Generate public URL
                        url = f"https://{self.s3_bucket}.s3.amazonaws.com/{key}"
                        upload_urls[format_type] = url

                        logger.info(f"✅ Uploaded {format_type} ({os.path.getsize(local_path):,} bytes): {url}")

                    except Exception as e:
                        logger.warning(f"❌ Failed to upload {format_type} to S3: {e}")

            if upload_urls:
                logger.info(f"📁 All files uploaded to conversation folder: conversations/{conversation_id}/reports/")

        except Exception as e:
            logger.warning(f"S3 upload failed (AWS credentials may not be configured): {e}")
            logger.warning("Files are still available locally for download")

        return upload_urls
